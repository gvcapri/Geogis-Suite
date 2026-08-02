import os
import re
import pandas as pd
from docx import Document
import unicodedata

# ==========================
# CONFIGURAÇÕES
# ==========================
ARQUIVO_WORD_CRF = r"CRF - JARDIM DE ALÁ_ATT.docx"
ARQUIVO_EXCEL_BASE = r"MET_Jd_Ala_VG.xlsx"
ARQUIVO_SAIDA = "comparativo_perimetro_area.xlsx"

# ==========================
# 1. FUNÇÕES DE LIMPEZA (HÍBRIDA)
# ==========================

def normalizacao_blindada(val):
    """
    Lógica Híbrida:
    1. Se tiver LETRAS (ex: Area Remanescente, Lote 10A): Juntas tudo (AREAREMANESCENTE1).
    2. Se for SÓ NÚMEROS compostos (ex: 02/03): Separa com underline (2_3) para não confundir com 23.
    """
    if pd.isna(val) or str(val).strip() == "": return ""
    
    # Remove acentos e converte para maiúsculo
    s = str(val).upper().strip()
    s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
    
    # Remove palavras que não fazem parte do ID
    s_limpa = s.replace("QUADRA", "").replace("LOTE", "").replace("INDICAÇÃO", "").replace("NUMERICA", "")
    
    # Verifica se sobraram LETRAS (A-Z) no identificador
    tem_letras = bool(re.search(r'[A-Z]', s_limpa))
    
    if tem_letras:
        # --- CASO 1: TEM LETRAS (Áreas, Lote A, etc) ---
        # Comportamento antigo: Concatena tudo
        # Ex: "AREA REMANESCENTE 01" -> "AREAREMANESCENTE1"
        partes = re.findall(r'[A-Z0-9]+', s_limpa)
        resultado = []
        for p in partes:
            if p.isdigit():
                resultado.append(str(int(p))) # Tira zero à esquerda
            else:
                resultado.append(p)
        return "".join(resultado)
        
    else:
        # --- CASO 2: APENAS NÚMEROS (Lotes compostos) ---
        # Comportamento novo: Usa underline para evitar colisão 2/3 vs 23
        # Ex: "02/03" -> "2_3"
        partes = re.split(r'[ \/\-\\&]+', s_limpa)
        resultado = []
        for p in partes:
            p_num = re.sub(r'[^0-9]', '', p)
            if p_num:
                resultado.append(str(int(p_num))) # Tira zero à esquerda
        
        return "_".join(resultado)

def converter_para_float(valor):
    if pd.isna(valor) or str(valor).strip() == "": return 0.0
    s = str(valor).strip().lower().replace("m²", "").replace("m", "").strip()
    if "," in s: s = s.replace(".", "").replace(",", ".")
    try:
        s_limpa = re.sub(r"[^\d.]", "", s)
        return round(float(s_limpa), 2)
    except:
        return 0.0

# ==========================
# 2. EXTRAÇÃO DO WORD
# ==========================

def ler_word_crf(caminho):
    if not os.path.exists(caminho):
        print(f"❌ Arquivo Word não encontrado: {caminho}")
        return []

    print("📖 Lendo Word...")
    doc = Document(caminho)
    texto_completo = "\n".join([p.text for p in doc.paragraphs])
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                texto_completo += "\n" + c.text

    padrao_cabecalho = re.compile(
        r"INDICAÇÃO\s+NUMÉRICA:.*?QUADRA\s+(?P<quadra>.*?)\s+LOTE\s+(?P<lote>.*?)(?=\n|QUALIFICAÇÃO|MODALIDADE)",
        re.IGNORECASE | re.DOTALL
    )

    resultados = []
    matches = list(padrao_cabecalho.finditer(texto_completo))
    
    for i, m in enumerate(matches):
        q_vis = m.group("quadra").strip()
        l_vis = m.group("lote").strip()
        
        inicio = m.end()
        fim = matches[i+1].start() if i + 1 < len(matches) else len(texto_completo)
        texto_bloco = texto_completo[inicio:fim]
        
        m_peri = re.search(r"PERÍMETRO:\s*([\d\.,]+)", texto_bloco, re.IGNORECASE)
        m_area = re.search(r"ÁREA(?: TOTAL)?:\s*([\d\.,]+)", texto_bloco, re.IGNORECASE)
        
        resultados.append({
            "Quadra_Word": q_vis,
            "Lote_Word": l_vis,
            "ID_JOIN": f"{normalizacao_blindada(q_vis)}_{normalizacao_blindada(l_vis)}",
            "Peri_Word": converter_para_float(m_peri.group(1)) if m_peri else 0.0,
            "Area_Word": converter_para_float(m_area.group(1)) if m_area else 0.0
        })
        
    return resultados

# ==========================
# 3. PROCESSAMENTO
# ==========================

# A. Ler Word
dados_word = ler_word_crf(ARQUIVO_WORD_CRF)
df_word = pd.DataFrame(dados_word)

# B. Ler Excel
print("📊 Lendo Excel...")
if os.path.exists(ARQUIVO_EXCEL_BASE):
    df_ex_raw = pd.read_excel(ARQUIVO_EXCEL_BASE)
    
    col_q = next((c for c in df_ex_raw.columns if "Quadra" in str(c)), df_ex_raw.columns[0])
    col_l = next((c for c in df_ex_raw.columns if "Lote" in str(c)), df_ex_raw.columns[1])
    col_p = next((c for c in df_ex_raw.columns if "Perímetro" in str(c)), None)
    col_a = next((c for c in df_ex_raw.columns if "rea" in str(c)), None)

    df_excel = pd.DataFrame()
    # Aplica a mesma normalização híbrida
    df_excel["ID_JOIN"] = df_ex_raw.apply(
        lambda x: f"{normalizacao_blindada(x[col_q])}_{normalizacao_blindada(x[col_l])}", axis=1
    )
    
    df_excel["Peri_Excel"] = df_ex_raw[col_p].apply(converter_para_float) if col_p else 0.0
    df_excel["Area_Excel"] = df_ex_raw[col_a].apply(converter_para_float) if col_a else 0.0
    df_excel["Quadra_Ex"] = df_ex_raw[col_q]
    df_excel["Lote_Ex"] = df_ex_raw[col_l]
else:
    print("❌ Excel não encontrado.")
    exit()

# C. Cruzamento
print("🔄 Cruzando dados...")
df_final = pd.merge(df_word, df_excel, on="ID_JOIN", how="outer")

# Preenchimento Visual
df_final["Quadra"] = df_final["Quadra_Word"].fillna(df_final["Quadra_Ex"])
df_final["Lote"] = df_final["Lote_Word"].fillna(df_final["Lote_Ex"])

# Cálculos
df_final["Dif_Peri"] = (df_final["Peri_Word"].fillna(0) - df_final["Peri_Excel"].fillna(0)).round(2)
df_final["Dif_Area"] = (df_final["Area_Word"].fillna(0) - df_final["Area_Excel"].fillna(0)).round(2)

# Status
def definir_status(row):
    erros = []
    if pd.isna(row["Quadra"]): return ""

    if row["Peri_Word"] == 0 and row["Area_Word"] == 0: return "Não encontrado no Word"
    if pd.isna(row["Peri_Excel"]): return "Não encontrado no Excel"
    
    if abs(row["Dif_Peri"]) > 0.05: erros.append("Perímetro")
    if abs(row["Dif_Area"]) > 0.05: erros.append("Área")
    
    if not erros: return "✅ OK"
    return f"🚩 Divergência: {' e '.join(erros)}"

df_final["Status"] = df_final.apply(definir_status, axis=1)

# Organização
colunas_finais = [
    "Quadra", "Lote", 
    "Area_Word", "Area_Excel", "Dif_Area",
    "Peri_Word", "Peri_Excel", "Dif_Peri", 
    "Status"
]
df_saida = df_final[colunas_finais].drop_duplicates()

# ==========================
# 4. ORDENAÇÃO
# ==========================
print("🔢 Ordenando...")
def chave_sort(v):
    s = str(v).strip()
    nums = re.findall(r'\d+', s)
    return int(nums[0]) if nums else float('inf'), s

df_saida["_q"] = df_saida["Quadra"].apply(lambda x: chave_sort(x)[0])
df_saida["_l"] = df_saida["Lote"].apply(lambda x: chave_sort(x)[0])

df_saida = df_saida.sort_values(by=["_q", "Quadra", "_l", "Lote"]).drop(columns=["_q", "_l"])

# ==========================
# SALVAR
# ==========================
df_saida.to_excel(ARQUIVO_SAIDA, index=False)
print(f"✅ Concluído: {ARQUIVO_SAIDA}")