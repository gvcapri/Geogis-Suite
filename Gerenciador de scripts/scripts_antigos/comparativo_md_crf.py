import re
import pandas as pd
from docx import Document
import os

# =========================
# 1. FUNÇÕES UTILITÁRIAS
# =========================

def ler_docx(caminho):
    if not os.path.exists(caminho):
        print(f"❌ Erro: Arquivo não encontrado: {caminho}")
        return ""
    doc = Document(caminho)
    texto = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto += "\n" + cell.text
    return texto.upper()

def str_para_float(valor):
    if not valor or pd.isna(valor):
        return 0.0
    limpo = str(valor).lower().replace("m²", "").replace("m", "").strip()
    limpo = re.sub(r'[^0-9,.]', '', limpo) 
    if not limpo: return 0.0
    if "," in limpo and "." in limpo:
        limpo = limpo.replace(".", "").replace(",", ".")
    elif "," in limpo:
        limpo = limpo.replace(",", ".")
    try:
        return float(limpo)
    except ValueError:
        return 0.0

def normalizar_identificador(valor):
    """
    Normaliza mantendo nomes complexos como 'AREA VERDE' ou 'AREA REMANESCENTE'.
    """
    if not valor: return ""
    v = str(valor).strip().upper()
    
    # Remove apenas termos que são estritamente rótulos de campo
    v = v.replace("LOTE", "").replace("QUADRA", "").replace("INDICAÇÃO", "").replace("NUMÉRICA", "").strip()
    
    # Se for apenas número, limpa zeros à esquerda
    if v.isdigit():
        return v.lstrip('0') if v.lstrip('0') != "" else "0"
    
    # Para nomes compostos: remove caracteres especiais mas MANTÉM letras e espaços
    v = re.sub(r'[^0-9A-ZÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕ ]', '', v)
    v = re.sub(r'\s+', ' ', v).strip()
    
    # Normaliza números dentro do nome (ex: REMANESCENTE 02 -> REMANESCENTE 2)
    v = re.sub(r'\b0+(\d+)', r'\1', v)
    
    return v

# =========================
# 2. EXTRAÇÃO DE DADOS
# =========================

def extrair_md(texto):
    dados = []
    # Divide o texto em blocos para não misturar dados de lotes diferentes
    blocos = re.split(r"PROPRIEDADE:", texto)
    for bloco in blocos:
        if not bloco.strip(): continue
        # Regex captura o nome do lote completo, incluindo espaços e acentos
        m = re.search(r"LOTE:\s*([0-9A-Z\-\/ÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕ ]+)\s*-\s*QUADRA:\s*([0-9A-Z ]+)", bloco)
        if m:
            quadra = normalizar_identificador(m.group(2))
            lote = normalizar_identificador(m.group(1))
            
            area = re.search(r"ÁREA\s*(?:TOTAL)?\s*:?\s*([\d.,]+)", bloco)
            peri = re.search(r"PERÍMETRO\s*:?\s*([\d.,]+)", bloco)
            
            val_area = str_para_float(area.group(1)) if area else 0.0
            val_peri = str_para_float(peri.group(1)) if peri else 0.0
            
            if quadra and lote and (val_area > 0 or val_peri > 0):
                dados.append({"quadra": quadra, "lote": lote, "area_md": val_area, "perimetro_md": val_peri})
    return dados

def extrair_crf(texto):
    dados = []
    # Divide por indicação numérica para garantir que a área pertença à quadra/lote corretos
    blocos = re.split(r"INDICAÇÃO\s+NUMÉRICA:", texto)
    for bloco in blocos:
        if not bloco.strip(): continue
        m_q = re.search(r"QUADRA\s*([0-9A-Z]+)", bloco)
        # Captura o lote até o fim da linha para pegar nomes como 'AREA VERDE 01'
        m_l = re.search(r"LOTE\s*([0-9A-ZÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕ\s\/]+)(?:\n|$|\r)", bloco)
        
        if m_q and m_l:
            quadra = normalizar_identificador(m_q.group(1))
            lote = normalizar_identificador(m_l.group(1))
            
            area = re.search(r"ÁREA\s+TOTAL:\s*([\d.,]+)", bloco)
            peri = re.search(r"PERÍMETRO:\s*([\d.,]+)", bloco)
            
            val_area = str_para_float(area.group(1)) if area else 0.0
            val_peri = str_para_float(peri.group(1)) if peri else 0.0
            
            if quadra and lote:
                dados.append({"quadra": quadra, "lote": lote, "area_crf": val_area, "perimetro_crf": val_peri})
    return dados

# =========================
# 3. PROCESSAMENTO FINAL
# =========================

ARQUIVO_MD = r"MD_Lotes_Integral_JD_de_Alá.docx"
ARQUIVO_CRF = r"CRF - JARDIM DE ALÁ_ATT.docx"

print("⏳ Lendo e processando arquivos...")
df_md = pd.DataFrame(extrair_md(ler_docx(ARQUIVO_MD)))
df_crf = pd.DataFrame(extrair_crf(ler_docx(ARQUIVO_CRF)))

# Agrupar apenas para remover duplicatas reais de texto normalizado
df_md = df_md.groupby(['quadra', 'lote'], as_index=False).sum()
df_crf = df_crf.groupby(['quadra', 'lote'], as_index=False).sum()

# Cruzamento
df = pd.merge(df_md, df_crf, on=["quadra", "lote"], how="outer").fillna(0)

# Filtro de segurança contra linhas vazias
df = df[(df['area_md'] > 0) | (df['area_crf'] > 0)]

# Cálculos
df["dif_area"] = (df["area_md"] - df["area_crf"]).round(4)
df["dif_perimetro"] = (df["perimetro_md"] - df["perimetro_crf"]).round(4)

def definir_status(row):
    tol = 0.01
    # round(2) trata a margem de erro de 1cm conforme solicitado
    e_area = abs(round(row['dif_area'], 2)) > tol
    e_peri = abs(round(row['dif_perimetro'], 2)) > tol
    if e_area and e_peri: return "DIVERGÊNCIA: ÁREA E PERÍMETRO"
    if e_area: return "DIVERGÊNCIA: ÁREA"
    if e_peri: return "DIVERGÊNCIA: PERÍMETRO"
    return "OK"

df["STATUS"] = df.apply(definir_status, axis=1)

# Ordenação Humana
def human_sort(val):
    numeros = re.findall(r'\d+', str(val))
    return int(numeros[0]) if numeros else 9999

df = df.sort_values(by=['quadra', 'lote'], key=lambda x: x.map(human_sort))

df.to_excel("CONFERENCIA_FINAL_REVISADA.xlsx", index=False)
print("✅ Planilha gerada! Áreas Verdes e Remanescentes agora estão com nomes completos.")