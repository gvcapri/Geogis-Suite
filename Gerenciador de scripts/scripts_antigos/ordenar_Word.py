import re
from docx import Document

# Caminhos
docx_path = r"ARQUIVO BASE\md_final (54).docx"
output_path = "Ordenado - Cidade Alta_Juruena.docx"

# --------------------------
# Funções auxiliares
# --------------------------

# Extrai blocos de texto completos
def extrair_blocos_completos(doc_path):
    doc = Document(doc_path)
    blocos = {}
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        text = paragraph.text.upper().strip()
        if text.startswith("INDICAÇÃO NUMÉRICA:"):
            ident = text.replace("INDICAÇÃO NUMÉRICA:", "").strip()
            blocos[ident] = [paragraph]
            i += 1
            while i < len(paragraphs):
                next_text = paragraphs[i].text.upper().strip()
                if next_text.startswith("INDICAÇÃO NUMÉRICA:"):
                    break
                blocos[ident].append(paragraphs[i])
                i += 1
        else:
            i += 1
    return blocos

# Copia parágrafo com formatação
def copiar_paragrafo_formatado(par_source, par_destino):
    for run in par_source.runs:
        r = par_destino.add_run(run.text)
        r.bold = run.bold
        r.italic = run.italic
        r.underline = run.underline
        r.font.name = run.font.name
        r.font.size = run.font.size

# Extrai e normaliza quadra/lote
def extrair_chave(ident):
    m = re.match(r"QUADRA\s+(.+?)\s+LOTE\s+(.+)", ident.upper())
    if m:
        quadra = m.group(1).strip()
        lote = m.group(2).strip()
        return (quadra, lote)
    return ("", "")

# Função de ordenação
def chave_ordem_quadralote(ident):
    quadra, lote = extrair_chave(ident)

    def split_alphanum(text):
        return [
            part.zfill(10) if part.isdigit() else part.upper()
            for part in re.split(r'(\d+)', text) if part
        ]

    quadra_key = split_alphanum(quadra)

    # Lógica: nomes completos primeiro, depois curtos
    if lote.isalpha() or len(lote) > 4:
        lote_class = (0,)  # nomes completos
    else:
        lote_class = (1,)  # códigos curtos

    lote_key = split_alphanum(lote)
    return (quadra_key, lote_class, lote_key)

# --------------------------
# Execução
# --------------------------

# Extrai blocos
blocos = extrair_blocos_completos(docx_path)

# Ordena os identificadores com base nas chaves definidas
identificadores_ordenados = sorted(blocos.keys(), key=chave_ordem_quadralote)

# Cria novo documento
novo_doc = Document()

# Adiciona blocos ao novo documento na ordem correta
for ident in identificadores_ordenados:
    for p in blocos[ident]:
        par_novo = novo_doc.add_paragraph()
        copiar_paragrafo_formatado(p, par_novo)
    novo_doc.add_paragraph()  # quebra entre blocos

# Salva
novo_doc.save(output_path)
print(f"Documento salvo com sucesso: {output_path}")
