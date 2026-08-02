import re
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

ARQUIVO_ENTRADA = "c:\Users\Rafael Honorato\AppData\Local\Temp\tmp8BB0.vdml"
ARQUIVO_SAIDA = "MD_lotes.docx"


def texto_documento(bloco):
    return "\n".join(p.text.strip() for p in bloco if p.text.strip())


def extrair_chave_ordenacao(bloco):
    texto = texto_documento(bloco)

    # Quadra
    quadra_match = re.search(r"QUADRA[:\s\-]*([0-9]+)", texto, re.IGNORECASE)
    if not quadra_match:
        return (9999, 9999, 9, "")

    quadra = int(quadra_match.group(1))

    # Lote (captura texto bruto)
    lote_match = re.search(r"LOTE[:\s\-]*([0-9A-Za-z\.\-]+)", texto, re.IGNORECASE)
    if not lote_match:
        return (quadra, 9999, 9, "")

    lote_raw = lote_match.group(1)

    # Extrai TODOS os números do lote
    numeros = [int(n) for n in re.findall(r"\d+", lote_raw)]

    if not numeros:
        # lote sem número (caso raro)
        return (quadra, 9999, 9, lote_raw)

    if len(numeros) > 1:
        # lote composto → usar o MAIOR número
        tipo = 2
        num_base = max(numeros)
        sufixo = ""
    else:
        num_base = numeros[0]
        tipo = 0
        sufixo_match = re.search(rf"{num_base}([A-Za-z])", lote_raw)
        sufixo = sufixo_match.group(1) if sufixo_match else ""

    return (quadra, num_base, tipo, sufixo)



def remover_paginas_em_branco(doc):
    for p in list(doc.paragraphs):
        if not p.text.strip():
            el = p._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)


def copiar_bloco(dest_doc, bloco):
    for p in bloco:
        novo_p = dest_doc.add_paragraph()
        novo_p._element.clear()
        novo_p._element.extend(deepcopy(p._element))


# ================== LEITURA DO DOCUMENTO ==================

doc = Document(ARQUIVO_ENTRADA)

blocos = []
bloco_atual = []

for p in doc.paragraphs:
    bloco_atual.append(p)

    if "Datum" in p.text or "SIRGAS2000" in p.text:
        blocos.append(bloco_atual)
        bloco_atual = []

if bloco_atual:
    blocos.append(bloco_atual)

# ================== ORDENAÇÃO ==================

blocos_ordenados = sorted(blocos, key=extrair_chave_ordenacao)

# ================== ESCRITA DO NOVO DOCUMENTO ==================

doc_out = Document()

for i, bloco in enumerate(blocos_ordenados):
    copiar_bloco(doc_out, bloco)

    # quebra lógica após cada memorial (limitação visual aceita)
    if i < len(blocos_ordenados) - 1:
        doc_out.add_page_break()

# limpeza final
remover_paginas_em_branco(doc_out)

doc_out.save(ARQUIVO_SAIDA)

print(f"Arquivo gerado com sucesso: {ARQUIVO_SAIDA}")
