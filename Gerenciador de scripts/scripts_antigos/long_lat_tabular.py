import pandas as pd
import re
from docx import Document
from docx.shared import Pt

def natural_sort_key(s):
    """Função para ordenar corretamente números e letras."""
    return tuple(int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', str(s)))

print("1. Lendo e ordenando a planilha...")
df = pd.read_excel('met_saojorge.xlsx', dtype=str)
df.columns = df.columns.str.strip()

# Ordenar em ordem crescente
df['Quadra_sort'] = df['Quadra'].apply(natural_sort_key)
df['Lote_sort'] = df['Lote'].apply(natural_sort_key)
df = df.sort_values(by=['Quadra_sort', 'Lote_sort']).drop(columns=['Quadra_sort', 'Lote_sort'])

lotes_ordenados = df.to_dict('records')

print("2. Abrindo o arquivo Word original...")
try:
    doc = Document('md_lotes_croqui.docx')
except Exception as e:
    print(f"Erro ao abrir o Word. Feche o arquivo se estiver aberto! Erro: {e}")
    exit()

print("3. Preenchendo as tabelas de cada página...")
indice_lote = 0

# O script vai procurar em todas as tabelas do documento
for table in doc.tables:
    # Se já preencheu todos os lotes da planilha, ele para
    if indice_lote >= len(lotes_ordenados):
        break

    try:
        # Pega a primeira célula da tabela (canto superior esquerdo)
        celula = table.cell(0, 0)
        texto_celula = celula.text.strip()

        # Verifica se essa é a tabela de cabeçalho (que contém "Código da parcela")
        if "Código da parcela" in texto_celula:
            lote_atual = lotes_ordenados[indice_lote]
            quadra = str(lote_atual['Quadra']).strip()
            lote = str(lote_atual['Lote']).strip()

            # Limpa o texto original da célula
            celula.text = ""
            
            # Insere o novo texto com a Quadra e o Lote
            run = celula.paragraphs[0].add_run(f"Quadra {quadra} - Lote {lote}")
            
            # Define a fonte para ficar elegante e padronizada com o seu arquivo
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True

            print(f"✔ Preenchido na página: Quadra {quadra} - Lote {lote}")
            indice_lote += 1
            
    except Exception as e:
        # Ignora tabelas menores ou mal formatadas que não sejam o cabeçalho
        continue

print("4. Salvando o novo arquivo...")
arquivo_final = 'memorial_croqui_FINALIZADO.docx'
doc.save(arquivo_final)

print(f"\nSucesso! {indice_lote} páginas preenchidas e salvas como '{arquivo_final}'.")